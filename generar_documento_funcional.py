#!/usr/bin/env python3
"""Genera el documento funcional Word del proyecto CALCULO_BIOMASA."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def build_document(output_path: Path) -> None:
    doc = Document()

    title = doc.add_heading("Documento funcional — Reporte de biomasa (CALCULO_BIOMASA)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Fecha: {date.today().strftime('%d/%m/%Y')}\n").italic = True
    meta.add_run("Destinatarios: Equipo FOOD / Produccion / Control de gestion\n").italic = True
    meta.add_run("Estado premisas de negocio: Validadas (referencia abril 2026)").italic = True

    add_heading(doc, "1. Proposito del desarrollo", 1)
    add_para(
        doc,
        "CALCULO_BIOMASA es una herramienta de consulta y reporting que consolida, en un unico "
        "informe HTML interactivo, el seguimiento diario de la biomasa en planta: entradas, "
        "consumo en cajas (TINA), salidas de producto, desglose stock/merma, arrastre de stock "
        "sin procesar y cruce con Business Central (BC) a nivel de lote/caja.",
    )
    add_para(
        doc,
        "El objetivo es dar al equipo FOOD una vision unificada, trazable y exportable (Excel) "
        "para analizar el flujo de materia prima, conciliar produccion con ERP y detectar "
        "desviaciones entre Innova y BC.",
    )

    add_heading(doc, "2. Que problema resuelve para FOOD", 1)
    add_bullets(
        doc,
        [
            "Centraliza metricas que antes requerian consultas manuales en Innova y BC.",
            "Aplica reglas de negocio acordadas (pkpackaging, balance de masa) de forma consistente.",
            "Permite seguimiento diario del balance entradas − cajas y entradas − salidas.",
            "Calcula stock y merma: Entrada TINA = Salidas CAJA + Stock + Merma.",
            "Cruza salidas Innova con albaranes BC por codigo de lote (number = Lot No.).",
            "Compara kilos Innova vs kilos BC ([Kilos] en Item Ledger Entry) en lotes enlazados.",
            "Facilita exportacion a Excel para reuniones, auditorias y analisis ad hoc.",
        ],
    )

    add_heading(doc, "3. Fuentes de datos", 1)
    add_table(
        doc,
        ["Sistema", "Objeto / tabla", "Uso"],
        [
            ["Innova (SQL Server)", "dbo.proc_packs + dbo.proc_materials", "Entradas, salidas, stock/merma por regtime"],
            ["Innova", "dbo.proc_matxacts + dbo.proc_packs", "Kg cajas = consumo TINA; fecha diaria = regtime de la tina (pack = proc_packs.id)"],
            ["Innova", "dbo.vw_stolt", "Arrastre stock sin procesar (fdespesque)"],
            ["Business Central (Azure SQL)", "bc.[Item Ledger Entry]", "Ventas, ajustes +/−, stock E/G; campo [Kilos] / Quantity"],
            ["Business Central", "bc.[Sales Shipment Line]", "Pedido ([Order No.]) del albaran"],
            ["Business Central", "bc.[Conversion productos]", "Cod. bascula Innova → Cod. producto BC"],
        ],
    )

    add_heading(doc, "4. Reglas de negocio (premisas validadas)", 1)
    add_table(
        doc,
        ["Concepto", "Regla"],
        [
            ["Entrada de biomasa", "proc_materials.pkpackaging = 3"],
            ["Salida", "pkpackaging <> 3 o NULL"],
            ["Stock", "Inventario TINA al cierre: stock inicial + Entradas − TINA procesada"],
            ["Merma", "Entrada TINA − Salidas CAJA − Stock (no es proc_packs.rtype)"],
            ["TINA procesada (kg)", "proc_matxacts (xactpath=1, nombre con 'tina'); fecha = proc_packs.regtime de la TINA"],
            ["Balance masa", "Entrada TINA = Salidas CAJA + Stock + Merma"],
        ],
    )
    add_para(doc, "Documento tecnico de referencia: PREMISAS.md en el repositorio del proyecto.")

    add_heading(doc, "5. Metricas y formulas principales", 1)
    add_table(
        doc,
        ["Metrica", "Formula / significado"],
        [
            ["Merma (kg)", "Entradas TINA − Salidas CAJA − Stock"],
            ["Diferencia (kg)", "Entradas TINA − TINA procesada"],
            ["Balance E-S (kg)", "Entradas − Salidas"],
            ["Stock sin procesar", "Arrastre acumulado (Entradas − Cajas o vw_stolt segun modo)"],
            ["% diferencia", "Diferencia / Entradas × 100"],
            ["Stock final teorico (opcional)", "Stock inicial + Entradas − Cajas"],
            ["Ajuste conciliacion (opcional)", "Stock final teorico − Stock final fisico"],
        ],
    )

    add_heading(doc, "6. Contenido del reporte HTML", 1)
    add_heading(doc, "6.1 KPIs de cabecera", 2)
    add_bullets(
        doc,
        [
            "Entradas TINA, salidas CAJA, stock inventario y merma (balance de masa)",
            "TINA procesada (kg y movimientos)",
            "Diferencia y balance entradas − salidas",
            "Stock sin procesar a fin de periodo",
            "Cruce BC: lotes enlazados, kg Innova vs kg BC, con/sin pedido",
            "Validacion de stock fisico (si se informan parametros opcionales)",
        ],
    )

    add_heading(doc, "6.2 Graficas interactivas", 2)
    add_bullets(
        doc,
        [
            "Evolucion diaria: entradas, cajas y salidas",
            "Diferencia diaria y acumulada",
            "Composicion del periodo: stock, merma, salidas, cajas",
            "Stock vs merma diarios en entradas",
            "Balance entradas − salidas por dia",
            "Todas maximizables en pantalla completa",
        ],
    )

    add_heading(doc, "6.3 Pestanas del informe HTML", 2)
    add_bullets(
        doc,
        [
            "Introduccion, Resumen, Graficas, Detalle diario, Balance (stock/merma)",
            "Cruce BC por lote (number = Lot No.)",
            "Balance BC E/G: Inicial + Produccion (Salidas CAJA) − Primera salida "
            "(Produccion = alta stock por coincidencia lote Innova∩BC)",
            "Balance por tipo en cajas: Inicial + Type 2 − Type 1 − Type 3",
            "Stock inicial / Stock final BC E/G por producto",
            "Analisis ILE (1/2/3): validacion de ecuacion y Type 3 por usuario/dia/producto",
            "Materiales y Debug (SQL)",
            "Pie unico: nota VAP + nota ajustes negativos Type 3",
        ],
    )

    add_heading(doc, "6.4 Tablas exportables a Excel", 2)
    add_bullets(
        doc,
        [
            "Detalle diario de produccion",
            "Entradas, salidas, stock y merma",
            "Cruce Innova / BC por lote (number / Lot No.)",
            "Balances BC E/G (kg y cajas) y analisis ILE",
            "Top 15 materiales de entrada y de salida",
            "Exportacion global en un unico libro Excel",
        ],
    )

    add_heading(doc, "7. Cruce Innova — Business Central", 1)
    add_para(doc, "Clave de enlace:", bold=True)
    add_para(doc, "dbo.proc_packs.number (codigo de lote/caja) = bc.[Item Ledger Entry].[Lot No.]")
    add_para(doc, "Por cada dia de regtime Innova, el reporte muestra:", bold=False)
    add_bullets(
        doc,
        [
            "Lotes de salida Innova y lotes enlazados en BC",
            "Kg Innova enlazados (proc_packs.weight)",
            "Kg BC enlazados (campo [Kilos], valor absoluto)",
            "Diferencia kg (Innova − BC) en lotes enlazados",
            "Unidades y kg BC con pedido vs sin pedido ([Order No.] del albaran)",
        ],
    )

    add_heading(doc, "8. Valores de referencia — abril 2026", 1)
    add_para(
        doc,
        "Totales obtenidos con las premisas validadas (ultima referencia operativa). "
        "Sirven como control al regenerar informes.",
    )
    add_table(
        doc,
        ["Metrica", "Valor"],
        [
            ["Entradas TINA", "518.532,75 kg"],
            ["TINA procesada", "353.246,00 kg"],
            ["Salidas CAJA", "410.802,15 kg · 71.174 cajas"],
            ["Stock de tinas", "169.158,75 kg"],
            ["Merma", "−61.428,15 kg (−11,85 %)"],
            ["Cruce BC lotes", "66.384 / 71.174 (93,27 %)"],
            ["Balance BC E/G (check kg)", "−70,56 kg"],
            ["Balance cajas", "Ini 6.829 · Ent 74.352 · Ven 73.320 · Adj− 3.884 · Check −228"],
            ["Analisis Type 3", "3 usuarios · top ACZ · muchos Type 3 con Kilos=0"],
        ],
    )

    add_heading(doc, "9. Como se genera el reporte", 1)
    add_para(doc, "Comando basico (modo legacy, periodo en dd/mm/aaaa):", bold=True)
    p = doc.add_paragraph()
    p.style = "Intense Quote"
    p.add_run(
        "python generar_reporte_biomasa.py --start 01/04/2026 --end 30/04/2026"
    )
    add_para(doc, "Salida:", bold=True)
    add_bullets(
        doc,
        [
            "Reports/reporte_biomasa_YYYYMMDD_YYYYMMDD.html",
            "Resumen en consola con totales del periodo",
            "Log de error en Reports/ si falla la ejecucion",
        ],
    )
    add_para(doc, "Opciones relevantes:", bold=True)
    add_bullets(
        doc,
        [
            "--skip-bc: genera el informe sin consultar Business Central",
            "--data-source vw_stolt_despesque: alternativa por fecha de despesque",
            "--stock-inicial / --stock-final-fisico: validacion opcional de stock fisico",
            "Credenciales en .env (Innova + BC)",
        ],
    )

    add_heading(doc, "10. Limitaciones y consideraciones", 1)
    add_bullets(
        doc,
        [
            "El cruce BC depende de que el lote (number) exista en Item Ledger Entry como [Lot No.].",
            "Lotes sin enlace (~13 % en marzo 2026) pueden deberse a desfase de fechas o contabilizacion pendiente.",
            "El pedido BC se obtiene del albaran (Sales Shipment Line), no del pack Innova.",
            "Stock/merma describe entradas en Innova, no inventario fisico de almacen salvo validacion opcional.",
            "Cambios en maestros (pkpackaging) requieren actualizar PREMISAS.md y regenerar mes de control.",
        ],
    )

    add_heading(doc, "11. Mantenimiento y soporte", 1)
    add_bullets(
        doc,
        [
            "Premisas canonicas: PREMISAS.md",
            "Implementacion: generar_reporte_biomasa.py (constantes PREMISA_* y SQL_*)",
            "Ante cambio de reglas: validar totales de un mes de referencia antes de usar en FOOD",
        ],
    )

    add_heading(doc, "12. Scripts Python del proyecto", 1)
    add_para(
        doc,
        "El repositorio incluye tres scripts ejecutables en la raiz del proyecto. Las credenciales "
        "de Innova y Business Central pueden definirse en el archivo .env o, para Innova, en el "
        "almacen seguro del sistema (keyring) mediante --save-creds.",
    )

    add_heading(doc, "12.1 generar_reporte_biomasa.py", 2)
    add_para(
        doc,
        "Genera el reporte HTML interactivo de biomasa (Innova + cruce BC + exportacion Excel). "
        "Si no se indican --start y --end, solicita el rango de fechas de forma interactiva.",
    )
    add_table(
        doc,
        ["Parametro", "Obligatorio", "Default", "Descripcion"],
        [
            ["--start", "No*", "—", "Fecha inicio dd/mm/aaaa"],
            ["--end", "No*", "—", "Fecha fin dd/mm/aaaa"],
            ["--server", "No", "DB_SERVER", "Servidor SQL Server Innova"],
            ["--database", "No", "DB_NAME", "Base de datos Innova"],
            ["--user", "No", "DB_USER / keyring", "Usuario Innova"],
            ["--password", "No", "DB_PASSWORD / keyring", "Contrasena Innova"],
            ["--cred-target", "No", "biomasa_sql_innova", "Identificador keyring para credenciales"],
            ["--save-creds", "No", "—", "Guarda credenciales Innova en keyring"],
            ["--output", "No", "Reports/ auto", "Ruta del HTML de salida"],
            ["--title", "No", "Reporte de Biomasa", "Titulo del informe"],
            ["--stock-inicial", "No", "—", "Stock inicial del periodo (kg)"],
            ["--stock-final-fisico", "No", "—", "Stock final fisico medido (kg)"],
            ["--data-source", "No", "legacy", "legacy o vw_stolt_despesque"],
            ["--skip-bc", "No", "—", "Omite consulta a Business Central"],
            ["--bc-server", "No", "BC_SERVER", "Servidor Azure SQL BC"],
            ["--bc-database", "No", "BC_DATABASE", "Base de datos BC"],
            ["--bc-user", "No", "BC_USER", "Usuario BC"],
            ["--bc-password", "No", "BC_PASSWORD", "Contrasena BC"],
            ["--bc-timeout", "No", "600 (BC_TIMEOUT)", "Timeout consulta BC (segundos)"],
            ["--bc-login-timeout", "No", "60 (BC_LOGIN_TIMEOUT)", "Timeout login BC (segundos)"],
        ],
    )
    add_para(doc, "* Si faltan --start y/o --end, el script pide las fechas por consola.", bold=False)
    add_para(doc, "Ejemplo:", bold=True)
    p1 = doc.add_paragraph()
    p1.style = "Intense Quote"
    p1.add_run("python generar_reporte_biomasa.py --start 01/03/2026 --end 31/03/2026")

    add_heading(doc, "12.2 validar_dia_salidas.py", 2)
    add_para(
        doc,
        "Validacion diaria: compara salidas Innova filtradas por regtime del dia con ventas BC "
        "enlazadas por lote (proc_packs.number = [Lot No.]). Carga ventas BC del mes completo "
        "y filtra por los lotes del dia. Las credenciales Innova se obtienen de .env o keyring.",
    )
    add_table(
        doc,
        ["Parametro", "Obligatorio", "Default", "Descripcion"],
        [
            ["--fecha", "Si", "—", "Fecha salida Innova dd/mm/aaaa"],
            ["--max-detalle", "No", "30", "Filas de detalle por lote en consola"],
            ["--bc-server", "No", "BC_SERVER", "Servidor Azure SQL BC"],
            ["--bc-database", "No", "BC_DATABASE", "Base de datos BC"],
            ["--bc-user", "No", "BC_USER", "Usuario BC"],
            ["--bc-password", "No", "BC_PASSWORD", "Contrasena BC"],
        ],
    )
    add_para(doc, "Ejemplo:", bold=True)
    p2 = doc.add_paragraph()
    p2.style = "Intense Quote"
    p2.add_run("python validar_dia_salidas.py --fecha 04/03/2026 --max-detalle 30")

    add_heading(doc, "12.3 generar_documento_funcional.py", 2)
    add_para(
        doc,
        "Regenera este documento Word (DOCUMENTO_FUNCIONAL_BIOMASA.docx). No acepta parametros "
        "de linea de comandos; la salida se escribe siempre en la raiz del proyecto.",
    )
    add_para(doc, "Ejemplo:", bold=True)
    p3 = doc.add_paragraph()
    p3.style = "Intense Quote"
    p3.add_run("python generar_documento_funcional.py")

    add_heading(doc, "12.4 Variables de entorno (.env)", 2)
    add_table(
        doc,
        ["Variable", "Usada por"],
        [
            ["DB_SERVER, DB_NAME, DB_USER, DB_PASSWORD", "Innova (reporte y validacion diaria)"],
            ["BC_SERVER, BC_DATABASE, BC_USER, BC_PASSWORD", "Business Central"],
            ["BC_TIMEOUT, BC_LOGIN_TIMEOUT", "Timeouts consulta BC (reporte principal)"],
        ],
    )

    add_heading(doc, "13. Resumen para el equipo FOOD", 1)
    add_para(
        doc,
        "Esta herramienta no sustituye Innova ni BC, pero automatiza el calculo acordado de biomasa "
        "y ofrece un panel unico para monitorizar produccion, conciliar salidas con ERP y exportar "
        "datos para analisis. El informe HTML es autocontenido (graficas, tablas, premisas visibles) "
        "y puede compartirse con el equipo sin necesidad de acceso directo a las bases de datos.",
    )

    doc.save(output_path)


if __name__ == "__main__":
    out = Path(__file__).resolve().parent / "DOCUMENTO_FUNCIONAL_BIOMASA.docx"
    build_document(out)
    print(f"Documento generado: {out}")
